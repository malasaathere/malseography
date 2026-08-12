from pathlib import Path
from docx import Document
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_TAB_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor

ROOT = Path(__file__).resolve().parents[1]
OUT_DIR = ROOT / "cv"
OUT_DIR.mkdir(exist_ok=True)
OUT_FILE = OUT_DIR / "Maleesha_Rajasooriya_Design_CV.docx"

# Preset: compact_reference_guide.
# Named override "design_cv_one_page": 0.55/0.66-inch margins,
# Arial 9.2 pt body, compact spacing, monochrome editorial styling.
PAGE_W_DXA, LR_DXA = 12240, 950
CONTENT_W_DXA = PAGE_W_DXA - 2 * LR_DXA
INK, MID, PALE, WHITE = "111318", "555B64", "EEF0F3", "FFFFFF"


def shade(cell, fill):
    pr = cell._tc.get_or_add_tcPr()
    node = pr.find(qn("w:shd")) or OxmlElement("w:shd")
    node.set(qn("w:fill"), fill)
    if node.getparent() is None:
        pr.append(node)


def cell_margins(cell, top=70, start=100, bottom=70, end=100):
    pr = cell._tc.get_or_add_tcPr()
    mar = pr.first_child_found_in("w:tcMar") or OxmlElement("w:tcMar")
    if mar.getparent() is None:
        pr.append(mar)
    for edge, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = mar.find(qn(f"w:{edge}")) or OxmlElement(f"w:{edge}")
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")
        if node.getparent() is None:
            mar.append(node)


def table_geometry(table, widths):
    table.autofit = False
    pr = table._tbl.tblPr
    tbl_w = pr.find(qn("w:tblW")) or OxmlElement("w:tblW")
    tbl_w.set(qn("w:w"), str(sum(widths)))
    tbl_w.set(qn("w:type"), "dxa")
    if tbl_w.getparent() is None:
        pr.append(tbl_w)
    ind = pr.find(qn("w:tblInd")) or OxmlElement("w:tblInd")
    ind.set(qn("w:w"), "0")
    ind.set(qn("w:type"), "dxa")
    if ind.getparent() is None:
        pr.append(ind)
    grid = table._tbl.tblGrid
    for old in list(grid):
        grid.remove(old)
    for width in widths:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(width))
        grid.append(col)
    for row in table.rows:
        for cell, width in zip(row.cells, widths):
            cell.width = Inches(width / 1440)
            tcw = cell._tc.get_or_add_tcPr().find(qn("w:tcW")) or OxmlElement("w:tcW")
            tcw.set(qn("w:w"), str(width))
            tcw.set(qn("w:type"), "dxa")
            if tcw.getparent() is None:
                cell._tc.get_or_add_tcPr().append(tcw)


def keep(paragraph, next_one=False):
    pr = paragraph._p.get_or_add_pPr()
    pr.append(OxmlElement("w:keepLines"))
    if next_one:
        pr.append(OxmlElement("w:keepNext"))


def border_bottom(paragraph):
    pr = paragraph._p.get_or_add_pPr()
    borders = pr.find(qn("w:pBdr")) or OxmlElement("w:pBdr")
    if borders.getparent() is None:
        pr.append(borders)
    bottom = OxmlElement("w:bottom")
    for key, val in (("val", "single"), ("sz", "8"), ("space", "4"), ("color", "D7DADE")):
        bottom.set(qn(f"w:{key}"), val)
    borders.append(bottom)


def hyperlink(paragraph, text, url):
    rid = paragraph.part.relate_to(url, "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink", is_external=True)
    link = OxmlElement("w:hyperlink")
    link.set(qn("r:id"), rid)
    run, rpr = OxmlElement("w:r"), OxmlElement("w:rPr")
    color, underline = OxmlElement("w:color"), OxmlElement("w:u")
    color.set(qn("w:val"), MID)
    underline.set(qn("w:val"), "single")
    rpr.extend([color, underline])
    run.append(rpr)
    node = OxmlElement("w:t")
    node.text = text
    run.append(node)
    link.append(run)
    paragraph._p.append(link)


def heading(doc, text):
    p = doc.add_paragraph(style="Heading 1")
    p.add_run(text.upper())
    keep(p, True)
    border_bottom(p)


def role(doc, title, org, dates, bullets):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(3)
    p.paragraph_format.space_after = Pt(0.5)
    p.paragraph_format.tab_stops.add_tab_stop(Inches(CONTENT_W_DXA / 1440), WD_TAB_ALIGNMENT.RIGHT)
    r = p.add_run(title)
    r.bold = True
    r.font.size = Pt(10)
    r.font.color.rgb = RGBColor.from_string(INK)
    p.add_run(f"  |  {org}").font.color.rgb = RGBColor.from_string(MID)
    r = p.add_run(f"\t{dates}")
    r.bold = True
    r.font.size = Pt(8.3)
    keep(p, True)
    for text in bullets:
        bp = doc.add_paragraph(style="CV Bullet")
        bp.add_run("• " + text)


doc = Document()
sec = doc.sections[0]
sec.page_width, sec.page_height = Inches(8.5), Inches(11)
sec.top_margin, sec.bottom_margin = Inches(0.55), Inches(0.50)
sec.left_margin = sec.right_margin = Inches(LR_DXA / 1440)
sec.header_distance, sec.footer_distance = Inches(0.2), Inches(0.2)

normal = doc.styles["Normal"]
normal.font.name, normal.font.size = "Arial", Pt(9.2)
normal.font.color.rgb = RGBColor.from_string(MID)
normal.paragraph_format.space_after, normal.paragraph_format.line_spacing = Pt(2.2), 1.03
h1 = doc.styles["Heading 1"]
h1.font.name, h1.font.size, h1.font.bold = "Arial", Pt(10.2), True
h1.font.color.rgb = RGBColor.from_string(INK)
h1.paragraph_format.space_before, h1.paragraph_format.space_after = Pt(7), Pt(3)
h1.paragraph_format.keep_with_next = True
bullet = doc.styles.add_style("CV Bullet", 1)
bullet.base_style = normal
bullet.font.name, bullet.font.size = "Arial", Pt(8.8)
bullet.paragraph_format.left_indent = Inches(0.17)
bullet.paragraph_format.first_line_indent = Inches(-0.17)
bullet.paragraph_format.space_after = Pt(1.2)
bullet.paragraph_format.keep_together = True

header = doc.add_table(rows=1, cols=2)
table_geometry(header, [7400, CONTENT_W_DXA - 7400])
header.style = "Table Grid"
for cell in header.rows[0].cells:
    shade(cell, INK)
    cell_margins(cell, 130, 170, 125, 170)
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    cell.paragraphs[0].paragraph_format.space_after = Pt(0)
left, right = header.rows[0].cells
p = left.paragraphs[0]
r = p.add_run("MALEESHA RAJASOORIYA")
r.bold, r.font.name, r.font.size = True, "Arial", Pt(24)
r.font.color.rgb = RGBColor.from_string(WHITE)
r = p.add_run("\nMOTION & VISUAL DESIGNER")
r.bold, r.font.size = True, Pt(9.2)
r.font.color.rgb = RGBColor.from_string("B8BDC5")
p = right.paragraphs[0]
p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
r = p.add_run("MALSEOGRAPHY")
r.bold, r.font.size = True, Pt(10.5)
r.font.color.rgb = RGBColor.from_string(WHITE)
r = p.add_run("\nDESIGN · MOTION · DIGITAL")
r.font.size = Pt(7.6)
r.font.color.rgb = RGBColor.from_string("B8BDC5")

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.paragraph_format.space_before, p.paragraph_format.space_after = Pt(4), Pt(3)
p.add_run("Sri Lanka   •   ")
hyperlink(p, "smrajasooriya@gmail.com", "mailto:smrajasooriya@gmail.com")
p.add_run("   •   ")
hyperlink(p, "Portfolio", "https://malasaathere.github.io")
p.add_run("   •   ")
hyperlink(p, "GitHub", "https://github.com/malasaathere")

heading(doc, "Profile")
doc.add_paragraph("Multidisciplinary designer and visual storyteller creating cinematic motion, social content, brand systems, and human-centred digital experiences. Combines strong typography and composition with structured Figma workflows, practical video production, and creative problem-solving.")

heading(doc, "Creative toolkit")
skills = doc.add_table(rows=2, cols=2)
table_geometry(skills, [CONTENT_W_DXA // 2, CONTENT_W_DXA - CONTENT_W_DXA // 2])
skills.style = "Table Grid"
skill_data = [
    ("MOTION & VIDEO", "After Effects — working proficiency · Premiere Pro — advanced\nMotion principles · pacing · transitions · titles · colour · sound"),
    ("DESIGN & PROTOTYPING", "Figma — advanced workflow · Photoshop — working knowledge · Illustrator — basic\nWireframes · prototypes · design systems · social assets"),
    ("VISUAL DIRECTION", "Typography · composition · hierarchy · visual identity\nCinematic storytelling · photography · editorial layouts"),
    ("WAYS OF WORKING", "Trend-aware concept development · creative problem-solving\nTeam coordination · feedback iteration · production handoff"),
]
for i, cell in enumerate(c for row in skills.rows for c in row.cells):
    cell_margins(cell, 80, 115, 75, 115)
    if i in (0, 3):
        shade(cell, PALE)
    title, body = skill_data[i]
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(1.5)
    r = p.add_run(title)
    r.bold, r.font.size = True, Pt(8.5)
    r.font.color.rgb = RGBColor.from_string(INK)
    p = cell.add_paragraph(body)
    p.paragraph_format.space_after, p.paragraph_format.line_spacing = Pt(0), 1.0
    p.runs[0].font.size = Pt(8.4)

heading(doc, "Relevant experience")
role(doc, "Assistant Media Director", "IMSSA, University of Kelaniya", "2026–Present", [
    "Lead visual identity, digital campaigns, and cinematic storytelling; coordinate media teams and guide creative direction.",
    "Translate event goals into concepts, social assets, motion content, and consistent cross-channel visuals.",
])
role(doc, "UI/UX Designer", "IdeaSprint & Independent Projects", "2025–2026", [
    "Led the Figma workflow and interactive prototyping for CeyLink, a finalist at the IdeaSprint 2025 intra-departmental hackathon.",
])
role(doc, "Level 1 Media Coordinator", "IMSSA, University of Kelaniya", "2025–2026", [
    "Produced digital content and visual communication for student events, adapting layouts and edits to fast-moving campaign needs.",
])
role(doc, "Freelance Photographer & Videographer", "Independent Creative", "2024–Present", [
    "Create high-contrast portraiture, event coverage, and promotional videos from concept through edit, colour, sound, and delivery.",
])

heading(doc, "Selected design projects")
projects = doc.add_table(rows=1, cols=3)
w = CONTENT_W_DXA // 3
table_geometry(projects, [w, w, CONTENT_W_DXA - 2 * w])
projects.style = "Table Grid"
project_data = [
    ("HACKX 11.0", "MOTION & CAMPAIGN DESIGN", "Directed a startup-challenge visual narrative spanning a 3D mascot, social layouts, motion graphics, and cinematic teaser scripts."),
    ("LEGACY IN BLOOM", "FIGMA · UX · DESIGN SYSTEM", "Designed an immersive responsive experience using research, interactive prototypes, custom visual motifs, and accessible structure."),
    ("CEYLON", "IDENTITY · TYPE · EDITORIAL", "Built a premium identity balancing cultural references with modern geometry, disciplined typography, packaging, and editorial layouts."),
]
for cell, (name, label, body) in zip(projects.rows[0].cells, project_data):
    cell_margins(cell, 90, 105, 85, 105)
    p = cell.paragraphs[0]
    r = p.add_run(name)
    r.bold, r.font.size = True, Pt(9)
    r.font.color.rgb = RGBColor.from_string(INK)
    p.paragraph_format.space_after = Pt(1)
    p = cell.add_paragraph(label)
    p.paragraph_format.space_after = Pt(2)
    p.runs[0].bold, p.runs[0].font.size = True, Pt(7.4)
    p = cell.add_paragraph(body)
    p.paragraph_format.space_after, p.paragraph_format.line_spacing = Pt(0), 1.0
    p.runs[0].font.size = Pt(8.1)

heading(doc, "Education")
p = doc.add_paragraph()
p.paragraph_format.space_after = Pt(0)
r = p.add_run("Undergraduate — Department of Industrial Management")
r.bold, r.font.size = True, Pt(9.4)
r.font.color.rgb = RGBColor.from_string(INK)
p.add_run("  |  University of Kelaniya  |  2025–Present")

footer = sec.footer.paragraphs[0]
footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = footer.add_run("PORTFOLIO AVAILABLE AT MALASAATHERE.GITHUB.IO")
r.bold, r.font.name, r.font.size = True, "Arial", Pt(7)
r.font.color.rgb = RGBColor.from_string(MID)

for table in doc.tables:
    for row in table.rows:
        row._tr.get_or_add_trPr().append(OxmlElement("w:cantSplit"))

doc.core_properties.title = "Maleesha Rajasooriya — Motion & Visual Designer CV"
doc.core_properties.subject = "Design-focused curriculum vitae"
doc.core_properties.author = "Maleesha Rajasooriya"
doc.save(OUT_FILE)
print(OUT_FILE)
